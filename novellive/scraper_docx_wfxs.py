import re
import time
import random
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


BASE_URL_TEMPLATE = "https://m.wfxs.tw/xiaoshuo/1203844/{chapter_id}/"
FIRST_CHAPTER_ID = 78804028

CHAPTER_RANGES = [
    (1, 60),
    (61, 120),
    (121, 180),
    (181, 211),
    (212, 270),
    (271, 300),
    (301, 360),
    (361, 421),
    (422, 481),
]

DOC_TITLE_PREFIX = "Novel"
OUTPUT_DIR = "docx_capitulos"

MIN_DELAY_SECONDS = 1.5
MAX_DELAY_SECONDS = 3.5

# tentativas por capítulo
MAX_RETRIES = 3
RETRY_BACKOFF_BASE = 2.0

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )
}


def chapter_number_to_id(chapter_number: int) -> int:
    return FIRST_CHAPTER_ID + (chapter_number - 1)


def chapter_url(chapter_number: int) -> str:
    return BASE_URL_TEMPLATE.format(chapter_id=chapter_number_to_id(chapter_number))


def clean_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def polite_delay(min_seconds: float = MIN_DELAY_SECONDS, max_seconds: float = MAX_DELAY_SECONDS) -> None:
    delay = random.uniform(min_seconds, max_seconds)
    print(f"  Aguardando {delay:.2f}s...")
    time.sleep(delay)


def retry_delay(attempt_number: int) -> None:
    delay = RETRY_BACKOFF_BASE ** attempt_number + random.uniform(0.3, 1.2)
    print(f"  Retry em {delay:.2f}s...")
    time.sleep(delay)


def fetch_chapter_paragraphs(session: requests.Session, chapter_number: int) -> tuple[str, list[str]]:
    url = chapter_url(chapter_number)
    response = session.get(url, headers=HEADERS, timeout=30)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "lxml")

    content_div = soup.find("div", id="read_conent_box", class_="entry")
    if content_div is None:
        raise ValueError(
            f"Não encontrei <div class='entry' id='read_conent_box'> no capítulo {chapter_number}: {url}"
        )

    title = None
    h1_title = soup.find("h1", class_="title")
    if h1_title:
        title = clean_text(h1_title.get_text(" ", strip=True))

    if not title:
        title = f"Capítulo {chapter_number}"

    paragraphs = []

    for p in content_div.find_all("p"):
        text = clean_text(p.get_text(" ", strip=True))
        if text:
            paragraphs.append(text)

    if not paragraphs:
        raw_text = content_div.get_text("\n", strip=True)
        for line in raw_text.splitlines():
            text = clean_text(line)
            if text:
                paragraphs.append(text)

    if not paragraphs:
        raise ValueError(f"Conteúdo vazio no capítulo {chapter_number}: {url}")

    return title, paragraphs


def fetch_chapter_with_retry(session: requests.Session, chapter_number: int) -> dict | None:
    last_error = None

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            title, paragraphs = fetch_chapter_paragraphs(session, chapter_number)

            if not paragraphs:
                print(f"  AVISO - capítulo {chapter_number} sem parágrafos")
                return None

            return {
                "number": chapter_number,
                "url": chapter_url(chapter_number),
                "title": title,
                "paragraphs": paragraphs,
            }

        except Exception as exc:
            last_error = exc
            print(f"  Tentativa {attempt}/{MAX_RETRIES} falhou no capítulo {chapter_number}: {exc}")

            if attempt < MAX_RETRIES:
                retry_delay(attempt)

    print(f"  ERRO FINAL - capítulo {chapter_number} falhou após {MAX_RETRIES} tentativas: {last_error}")
    return None


def configure_document_styles(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    heading_style = document.styles["Heading 1"]
    heading_style.font.name = "Times New Roman"
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True


def build_docx(chapters: list[dict], output_path: Path, range_start: int, range_end: int) -> None:
    document = Document()
    configure_document_styles(document)

    title = document.add_heading(f"{DOC_TITLE_PREFIX} Capítulos {range_start} a {range_end}", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    info = document.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info.add_run(f"Capítulos {range_start} a {range_end}")

    document.add_paragraph("")

    added_any = False

    for index, item in enumerate(chapters):
        print(f"Montando DOCX do capítulo {item['number']}")

        paragraphs = [str(p).strip() for p in item.get("paragraphs", []) if str(p).strip()]
        if not paragraphs:
            print(f"Capítulo {item.get('number')} sem parágrafos válidos, pulando.")
            continue

        heading = document.add_heading(item["title"], level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        document.add_paragraph("")
        document.add_paragraph("")

        for paragraph_text in paragraphs:
            p = document.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.add_run(paragraph_text)

        if index < len(chapters) - 1:
            document.add_page_break()

        added_any = True

    if not added_any:
        raise RuntimeError("Nenhum capítulo válido foi adicionado ao DOCX.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def fetch_range(session: requests.Session, range_start: int, range_end: int) -> tuple[list[dict], list[int]]:
    chapters_data = []
    failed_chapters = []

    for chapter_number in range(range_start, range_end + 1):
        print(f"Baixando capítulo {chapter_number}: {chapter_url(chapter_number)}")

        chapter_data = fetch_chapter_with_retry(session, chapter_number)
        if chapter_data is None:
            failed_chapters.append(chapter_number)
        else:
            chapters_data.append(chapter_data)
            print(
                f"  OK - {len(chapter_data['paragraphs'])} parágrafos capturados | "
                f"título: {chapter_data['title']}"
            )

        polite_delay()

    return chapters_data, failed_chapters


def main() -> None:
    session = requests.Session()
    output_dir = Path(OUTPUT_DIR)

    all_failed = {}

    for range_start, range_end in CHAPTER_RANGES:
        print(f"\n=== Processando capítulos {range_start} a {range_end} ===")
        chapters_data, failed_chapters = fetch_range(session, range_start, range_end)

        if not chapters_data:
            print(f"Nenhum capítulo válido encontrado para {range_start}-{range_end}.")
        else:
            output_file = output_dir / f"novel_capitulos_{range_start}_a_{range_end}.docx"
            build_docx(chapters_data, output_file, range_start, range_end)
            print(f"DOCX gerado com sucesso: {output_file.resolve()}")

        if failed_chapters:
            all_failed[f"{range_start}-{range_end}"] = failed_chapters

    if all_failed:
        print("\n=== RESUMO DE CAPÍTULOS QUE FALHARAM ===")
        for chapter_range, failures in all_failed.items():
            print(f"{chapter_range}: {failures}")
    else:
        print("\nTodos os capítulos foram processados com sucesso.")


if __name__ == "__main__":
    main()
